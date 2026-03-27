#!/usr/bin/env python3
"""
generate_word_cv.py

Generate main.docx – a Word version of the LaTeX CV (main.tex).

Design goals (matching the classicthesis / currvita LaTeX source)
──────────────────────────────────────────────────────────────────
• Font        : Palatino Linotype (closest Word equivalent of the
                Palatino / URW Palladio used by classicthesis)
• Font sizes  : Based on the 9 pt document base
                  \\normalsize  →  9.0 pt
                  \\large       → 10.95 pt
                  \\Large       → 12.0 pt
                  \\LARGE       → 14.4 pt
• Maroon      : #800000  (xcolor 'Maroon')
• Letter spacing
                  \\spacedallcaps      → +2.5 pt  (CV title)
                  \\spacedlowsmallcaps → +1.5 pt  (section headings)
• Section headings: Maroon bold first letter + small-caps letter-spaced rest
  preceded by a thin Maroon rule (matching classicthesis section dividers)
• Page margins: A4, total={7 in, 10.6 in}
                → side ≈ 0.635 in,  top/bottom ≈ 0.545 in
• Two-column borderless table for CV list items (3.2 cm date label | content)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Constants ────────────────────────────────────────────────────────────────
MAROON = RGBColor(0x80, 0x00, 0x00)
BLACK  = RGBColor(0x00, 0x00, 0x00)

FONT   = "Palatino Linotype"   # closest Word match for classicthesis Palatino

# Font sizes at 9 pt base (LaTeX size commands)
SZ     = 9.0     # \normalsize
SZ_LG  = 10.95   # \large
SZ_LG2 = 12.0    # \Large
SZ_LG3 = 14.4    # \LARGE

# Character spacing added by classicthesis tracking commands
SP_TITLE = Pt(2.5)   # \spacedallcaps      – main CV title
SP_SEC   = Pt(1.5)   # \spacedlowsmallcaps – section headings

# Page geometry  (LaTeX: a4paper, total={7in, 10.6in})
PAGE_W = Inches(8.27)     # A4 width
PAGE_H = Inches(11.69)    # A4 height
MAR_LR = Inches(0.635)    # left / right margin  = (8.27 - 7) / 2
MAR_TB = Inches(0.545)    # top  / bottom margin = (11.69 - 10.6) / 2
TW     = Inches(7)        # total text width

# Two-column layout for CV list items
LABEL_W   = Cm(3.2)
CONTENT_W = TW - LABEL_W


# ─── Low-level helpers ────────────────────────────────────────────────────────

def run(para, text, size=SZ, bold=False, italic=False, color=None,
        small_caps=False, all_caps=False, spacing=None):
    """Append a typographically controlled run to *para*."""
    r = para.add_run(text)
    r.font.name      = FONT
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.italic    = italic
    r.font.color.rgb = color if color is not None else BLACK
    if small_caps: r.font.small_caps = True
    if all_caps:   r.font.all_caps   = True
    if spacing is not None:
        # python-docx ≥ 1.0 does not expose font.spacing; set via OOXML directly.
        # <w:spacing w:val="N"/> uses 1/20 pt (twips): Pt(x) → x*20 twips
        rPr = r._r.get_or_add_rPr()
        sp_el = OxmlElement('w:spacing')
        sp_el.set(qn('w:val'), str(int(spacing / 635)))  # EMU → twips
        rPr.append(sp_el)
    return r


def fmt(para, before=0, after=4, align=WD_ALIGN_PARAGRAPH.LEFT,
        keep_next=False, line_spacing=None):
    """Apply paragraph-level formatting to *para*."""
    pf = para.paragraph_format
    pf.space_before   = Pt(before)
    pf.space_after    = Pt(after)
    pf.alignment      = align
    pf.keep_with_next = keep_next
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def _no_table_borders(table):
    """Remove all cell and table borders from *table*."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'none')
        el.set(qn('w:sz'),    '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        bdr.append(el)
    tblPr.append(bdr)


def _zero_table_cell_margins(table):
    """Set tblCellMar to zero so cell widths are used without added padding."""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tcMar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    '0')
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tblPr.append(tcMar)


def _zero_cell_margins(cell):
    """Zero the padding on a single cell (for cells that hold nested tables)."""
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'),    '0')
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _cell_width(cell, width):
    """Fix a table cell to an exact width (width is a Length in EMUs)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width / 635)))   # EMU → twips (1/20 pt)
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)


def _top_border(para, color='800000'):
    """Attach a thin top-border rule to *para* (Maroon by default)."""
    pPr = para._p.get_or_add_pPr()
    bdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),   'single')
    top.set(qn('w:sz'),    '4')       # 0.5 pt line weight
    top.set(qn('w:space'), '4')       # 4 pt space between rule and text
    top.set(qn('w:color'), color)
    bdr.append(top)
    pPr.append(bdr)


def _hyperlink(para, text, url, size=SZ):
    """Append a clickable hyperlink run to *para*."""
    r_id = para.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    rr   = OxmlElement('w:r')
    rPr  = OxmlElement('w:rPr')

    fonts = OxmlElement('w:rFonts')
    fonts.set(qn('w:ascii'), FONT)
    fonts.set(qn('w:hAnsi'), FONT)
    rPr.append(fonts)

    for tag, val in (('w:sz', str(int(size * 2))), ('w:szCs', str(int(size * 2)))):
        el = OxmlElement(tag)
        el.set(qn('w:val'), val)
        rPr.append(el)

    col = OxmlElement('w:color')
    col.set(qn('w:val'), '0563C1')     # standard hyperlink blue
    rPr.append(col)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    rr.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    rr.append(t)

    hl.append(rr)
    para._p.append(hl)


def _add_nested_table(doc, cell, nrows, ncols, col_width):
    """
    Create a *nrows* × *ncols* table with each column at *col_width* EMUs,
    insert it inside *cell*, and fix both the tblGrid and tblW so Word /
    LibreOffice honour the explicit column widths.
    """
    tbl      = doc.add_table(rows=nrows, cols=ncols)
    tbl_elem = tbl._tbl
    doc.element.body.remove(tbl_elem)   # detach from document body

    # Fix tblW to the total explicit table width
    tblPr       = tbl_elem.find(qn('w:tblPr'))
    tblW        = tblPr.find(qn('w:tblW'))
    total_twips = str(int(col_width * ncols / 635))
    tblW.set(qn('w:w'),    total_twips)
    tblW.set(qn('w:type'), 'dxa')

    # Fix tblGrid (column blueprint used by the layout engine)
    tblGrid = tbl_elem.find(qn('w:tblGrid'))
    if tblGrid is not None:
        for gc in tblGrid.findall(qn('w:gridCol')):
            gc.set(qn('w:w'), str(int(col_width / 635)))

    # Attach to cell before its mandatory closing <w:p>
    tc     = cell._tc
    paras  = tc.findall(qn('w:p'))
    last_p = paras[-1]
    tc.insert(list(tc).index(last_p), tbl_elem)
    return tbl


# ─── Mid-level building blocks ────────────────────────────────────────────────

def section_heading(doc, first_letter, rest):
    """
    CV section heading matching the LaTeX pattern:
        \\textcolor{Maroon}{X}\\spacedlowsmallcaps{...}

    Rendered as: thin Maroon rule above + Maroon bold first letter
                 + small-caps letter-spaced remainder.
    """
    p = doc.add_paragraph()
    fmt(p, before=10, after=4, keep_next=True)
    _top_border(p)
    run(p, first_letter, size=SZ_LG, color=MAROON, bold=True, spacing=SP_SEC)
    run(p, rest,         size=SZ_LG, small_caps=True,          spacing=SP_SEC)


def item_table(doc):
    """Return a 2-column borderless table for CV list items."""
    tbl = doc.add_table(rows=0, cols=2)
    _no_table_borders(tbl)
    return tbl


def add_row(tbl, label, fill_fn):
    """
    Append one CV list row to *tbl*.
    *label*   – string for the right-aligned left column (date / key)
    *fill_fn* – callable(cell) that fills the right content cell
    """
    row       = tbl.add_row()
    lc, rc    = row.cells
    _cell_width(lc, LABEL_W)
    _cell_width(rc, CONTENT_W)

    lp = lc.paragraphs[0]
    fmt(lp, before=2, after=2, align=WD_ALIGN_PARAGRAPH.RIGHT)
    if label:
        run(lp, label)

    fill_fn(rc)


def bullet(cell, text):
    """Append a hanging-indent bullet paragraph to *cell*."""
    p  = cell.add_paragraph()
    fmt(p, before=0, after=1)
    pf = p.paragraph_format
    pf.left_indent       = Cm(0.5)
    pf.first_line_indent = Cm(-0.4)
    run(p, "\u2022 ", size=SZ)   # • bullet character
    run(p, text,      size=SZ)
    return p


# ─── Document builder ─────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    sec               = doc.sections[0]
    sec.page_width    = PAGE_W
    sec.page_height   = PAGE_H
    sec.left_margin   = MAR_LR
    sec.right_margin  = MAR_LR
    sec.top_margin    = MAR_TB
    sec.bottom_margin = MAR_TB

    # ── Base style ────────────────────────────────────────────────────────────
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(SZ)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(4)

    # ══════════════════════════════════════════════════════════════════════════
    # HEADER
    # ══════════════════════════════════════════════════════════════════════════

    # ── CV title: \cvheadingfont{\LARGE\color{Maroon}} + \spacedallcaps{...} ──
    p = doc.add_paragraph()
    fmt(p, before=0, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "CURRICULUM VIT\u00C6",             # Æ = U+00C6
        size=SZ_LG3, color=MAROON, spacing=SP_TITLE)

    # ── Name: \begin{Large} Adrien T'Kint \end{Large} ─────────────────────────
    p = doc.add_paragraph()
    fmt(p, before=4, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Adrien T\u2019Kint", size=SZ_LG2)  # ' = U+2019 (typographic apostrophe)

    # ── Subtitle: \begin{large}\textit{...}\end{large} ────────────────────────
    p = doc.add_paragraph()
    fmt(p, before=0, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p,
        "Results-driven multilingual engineer bridging R&D expertise and "
        "management training.\nSocial, ambitious and eager to learn.",
        size=SZ_LG, italic=True)

    # ══════════════════════════════════════════════════════════════════════════
    # PERSONAL INFORMATION
    # ══════════════════════════════════════════════════════════════════════════
    section_heading(doc, 'P', 'ersonal information')
    tbl = item_table(doc)

    def fill_birth(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=2)
        run(p, "5th of April 1993")

    def fill_address(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=2)
        run(p, "Rue des Francs 11, 1040 Brussels")

    def fill_phone(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=2)
        run(p, "+32 491/91 91 40")

    def fill_email(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=2)
        _hyperlink(p, "adtkint@gmail.com", "mailto:adtkint@gmail.com")

    add_row(tbl, "Birth date", fill_birth)
    add_row(tbl, "Address",    fill_address)
    add_row(tbl, "Phone",      fill_phone)
    add_row(tbl, "E-mail",     fill_email)

    # ══════════════════════════════════════════════════════════════════════════
    # WORK EXPERIENCES
    # ══════════════════════════════════════════════════════════════════════════
    section_heading(doc, 'W', 'ork experiences')
    tbl = item_table(doc)

    def fill_perseus(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=1)
        run(p, "Main systems engineer, then Product lead & Engineering team lead "
               "at Perseus Biomics, Leuven", bold=True)

        p2 = c.add_paragraph()
        fmt(p2, before=0, after=1)
        run(p2, "Driving the technical success of the MetaMAP, the flagship "
                "automated microbiome analyzer of Perseus Biomics by:")

        for b in (
            "Coordinating and aligning the efforts of internal and external "
            "experts in mechanics, electronics, software, firmware, microscopy "
            "and biochemistry using Gitlab issues and milestones",

            "Identifying, monitoring and mitigating development risks through "
            "structured risk analysis, proofs of concept and quantitative "
            "performance comparisons",

            "Progressively assuming ownership of the agile execution of the "
            "strategic product roadmap",

            "Key-contributor to multiple technical aspects of the device, from "
            "its first conceptual design (e.g. microscope architecture, thermal "
            "management) to its development (e.g. scanning routine coding, motor "
            "drivers optimization) all the way to its acceptance testing "
            "(e.g. measurement of the optical performances, functional testing)",

            "Ensuring device safety and successful CE marking",
        ):
            bullet(c, b)

    add_row(tbl, "2021 \u2192 present", fill_perseus)   # → = U+2192

    def fill_lambda(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=1)
        run(p, "Systems engineer for optical instrumentation at Lambda-X, Nivelles",
            bold=True)

        p2 = c.add_paragraph()
        fmt(p2, before=0, after=1)
        run(p2, "Consisting in design review, components choice, integration, "
                "alignment and testing of optical systems like:")

        for b in (
            "Space: IR telescope for Earth observation & Optical control system "
            "for the study of particles agglomeration in microgravity",

            "Ophthalmic: Coating viewer dedicated for intra-ocular lenses",

            "Industrial: Image guided laser system for cell therapy production "
            "& Car door projector",
        ):
            bullet(c, b)

    add_row(tbl, "2018 \u2192 2021", fill_lambda)

    def fill_teacher(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=2)
        run(p, "Mathematics and sciences teacher", bold=True)
        run(p, " at Lyc\u00e9e de Berlaymont in Waterloo")   # é = U+00E9

    add_row(tbl, "2017 \u2192 2018", fill_teacher)

    # ══════════════════════════════════════════════════════════════════════════
    # EDUCATION
    # ══════════════════════════════════════════════════════════════════════════
    section_heading(doc, 'E', 'ducation')
    tbl = item_table(doc)

    def fill_solvay(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=1)
        run(p, "Complementary Master in Management, "
               "specialization in Finances, Cum laude", bold=True)
        p2 = c.add_paragraph()
        fmt(p2, before=0, after=2)
        run(p2, "At the Solvay Business School of the "
                "Universit\u00e9 Libre de Bruxelles")    # é

    def fill_photonics(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=1)
        run(p, "European Master of Science in Photonics, Summa cum laude",
            bold=True)

        p2 = c.add_paragraph()
        fmt(p2, before=0, after=1)
        run(p2, "Joint degree (Universiteit Gent\u2013Vrije Universiteit Brussel)")
        # – = U+2013 en-dash

        p3 = c.add_paragraph()
        fmt(p3, before=0, after=1)
        run(p3, "Erasmus+", bold=True)
        run(p3, " grant for a year at the Karlsruhe Institute of Technology")

        p4 = c.add_paragraph()
        fmt(p4, before=0, after=2)
        run(p4, "Master thesis", bold=True)
        run(p4, " Qualification of an optical 3D sensor for industrial "
                "applications", italic=True)
        run(p4, " at SICK AG")

    def fill_bachelor(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=1)
        run(p, "Bachelor of Science in Electronics and Information Technology "
               "Engineering, Cum laude", bold=True)
        p2 = c.add_paragraph()
        fmt(p2, before=0, after=2)
        run(p2, "At the Vrije Universiteit Brussel")

    add_row(tbl, "2016 \u2192 2019", fill_solvay)
    add_row(tbl, "2014 \u2192 2016", fill_photonics)
    add_row(tbl, "2011 \u2192 2014", fill_bachelor)

    # ══════════════════════════════════════════════════════════════════════════
    # SKILLS
    # ══════════════════════════════════════════════════════════════════════════
    section_heading(doc, 'S', 'kills')
    tbl = item_table(doc)

    def fill_languages(c):
        # Nested 2×2 table – mirrors the LaTeX \begin{tabular}{|p{5cm}|p{5cm}}
        # Zero *this cell's* own padding so the nested table fills it exactly
        _zero_cell_margins(c)
        half     = CONTENT_W / 2
        lang_tbl = _add_nested_table(doc, c, 2, 2, col_width=half)
        _no_table_borders(lang_tbl)
        _zero_table_cell_margins(lang_tbl)
        # Minimise the mandatory closing paragraph spacing
        fmt(c.paragraphs[0], before=0, after=0)

        data = [
            ("French: Mother tongue", "Dutch: Very fluent"),
            ("English: Fluent (B2)",  "German: Fluent (B1)"),
        ]
        for r_idx, (left_text, right_text) in enumerate(data):
            for c_idx, text in enumerate((left_text, right_text)):
                cell_p = lang_tbl.rows[r_idx].cells[c_idx].paragraphs[0]
                _cell_width(lang_tbl.rows[r_idx].cells[c_idx], half)
                fmt(cell_p, before=1, after=1)
                run(cell_p, text)

    def fill_computer(c):
        p = c.paragraphs[0]
        fmt(p, before=2, after=2)
        run(p, "Agentic AI, Python, Matlab, Git, Lumerical, Zemax, Assembler, "
               "Java, LTspice, Autodesk Inventor, LaTeX, Word, Excel\u2026")
        # … = U+2026

    add_row(tbl, "Languages",       fill_languages)
    add_row(tbl, "Computer skills", fill_computer)

    # ══════════════════════════════════════════════════════════════════════════
    # SOCIAL LIFE
    # ══════════════════════════════════════════════════════════════════════════
    section_heading(doc, 'S', 'ocial Life')
    tbl = item_table(doc)

    social = [
        ("2022 \u2192 2025",
         "Coordinator of more than twenty organizers for the spring-summer "
         "trainings of the orienteering sport club ASUB. "
         "Main organizer of multiple competitions."),
        ("2019 \u2192 2021",
         "Member in charge of the finances of the local group of a political "
         "party in Auderghem"),
        ("2017 \u2192 2018",
         "Member of the board of direction of the Jeugd Parlement Jeunesse"),
        ("2015",
         "Main organizer of the summer course for the Board of European "
         "Students of Technology"),
        ("2010 \u2192 2011",
         "Scout leader at F\u00e9d\u00e9ration des Scouts Baden-Powell de Belgique"),
        # é = U+00E9
    ]

    for date, text in social:
        def fill_social(c, t=text):   # default-arg captures loop variable
            p = c.paragraphs[0]
            fmt(p, before=2, after=2)
            run(p, t)
        add_row(tbl, date, fill_social)

    # ── Save ──────────────────────────────────────────────────────────────────
    out = "main.docx"
    doc.save(out)
    print(f"\u2713 {out} written successfully.")


if __name__ == "__main__":
    build()
